"""Document builder — assembles a python-docx Document from parsed AST + config.

Covers: cover page, table of contents, headings, paragraphs (with inline
runs), academic three-line tables, images, ordered/unordered lists,
code blocks, blockquotes, thematic breaks, and LaTeX math equations.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Pt

# ---- Fix: when cairocffi cannot find the system Cairo DLL (common on
#      Windows), it raises OSError instead of ImportError.  This means
#      ``rlPyCairo.gstate``'s ``except ImportError: import cairo`` fallback
#      never triggers.  We install a tiny import-hook that wraps the OSError
#      into an ImportError so the pycairo fallback works.
#      The fix is idempotent and does nothing when cairo DLL is available. ----
_HOOK_INSTALLED = False


def _ensure_cairocffi_import_error() -> None:
    global _HOOK_INSTALLED
    if _HOOK_INSTALLED:
        return
    _HOOK_INSTALLED = True

    # Build a meta_path finder that intercepts "cairocffi" imports.
    # We must keep a reference to the *original* meta_path so that we can
    # delegate the real import attempt without infinite recursion.
    _REAL_PATH = list(sys.meta_path)

    class _Finder:
        def find_module(self, fullname, path=None):
            return self if fullname == "cairocffi" else None

        def load_module(self, fullname):
            if fullname in sys.modules:
                return sys.modules[fullname]
            saved = sys.meta_path
            sys.meta_path = _REAL_PATH
            try:
                __import__(fullname)
            except OSError:
                raise ImportError(
                    "cairocffi cannot load Cairo DLL → falling back to pycairo"
                ) from None
            finally:
                sys.meta_path = saved
            return sys.modules[fullname]

    sys.meta_path.insert(0, _Finder())


_ensure_cairocffi_import_error()
del _ensure_cairocffi_import_error

from md2docx.config import (
    AppConfig,
    CodeStyle,
    PageNumberConfig,
    StyleConfig,
    TableStyle,
    TextStyle,
)
from md2docx.math_handler import add_inline_math, add_math_paragraph
from md2docx.styles import (
    add_cover_line,
    apply_font_to_run,
    apply_format_to_para,
    setup_heading_styles,
    setup_normal_style,
    setup_toc_style,
    format_table_content,
    insert_toc_field,
    make_three_line_table,
    set_para_left_border,
    set_run_font,
    _set_table_border,
    _is_math_run,
    SIZE_SANHAO,
    SIZE_WUHAO,
)


# =============================================================================
# Inline rendering
# =============================================================================


def _render_inline_run(
    para,
    run_data: dict[str, Any],
    styles: StyleConfig,
    md_path: Path = Path("."),
) -> None:
    """Render a single inline run dict into a paragraph.

    Recursively handles bold, italic, code, links, inline math, images, and
    plain text.
    """
    rtype = run_data.get("type", "text")

    if rtype == "text":
        text = run_data.get("text", "")
        if text:
            run = para.add_run(text)
            apply_font_to_run(run, styles.body)

    elif rtype == "bold":
        children = run_data.get("children", [])
        for child in children:
            run = para.add_run(_extract_run_text(child))
            apply_font_to_run(run, styles.body)
            run.bold = True

    elif rtype == "italic":
        children = run_data.get("children", [])
        for child in children:
            run = para.add_run(_extract_run_text(child))
            apply_font_to_run(run, styles.body)
            run.italic = True

    elif rtype == "code":
        text = run_data.get("text", "")
        if text:
            run = para.add_run(text)
            apply_font_to_run(run, styles.code)

    elif rtype == "inline_math":
        latex = run_data.get("text", "")
        if latex:
            # Try OMML first
            success = add_inline_math(para, latex)
            if not success:
                run = para.add_run(f" ${latex}$ ")
                run.font.name = "Cambria Math"
                run.italic = True

    elif rtype == "image":
        src = run_data.get("src", "")
        alt = run_data.get("alt", "")
        _render_inline_image(para, src, alt, md_path)

    elif rtype == "linebreak":
        run = para.add_run("\n")
    elif rtype == "softbreak":
        run = para.add_run(" ")


def _extract_run_text(run_data: dict[str, Any]) -> str:
    """Get plain text from a run dict (recursive)."""
    rtype = run_data.get("type", "text")
    if rtype == "text":
        return run_data.get("text", "")
    elif rtype in ("bold", "italic", "link"):
        return _extract_children_text(run_data.get("children", []))
    elif rtype == "code":
        return run_data.get("text", "")
    else:
        return run_data.get("text", "")


def _extract_children_text(children: list[dict[str, Any]]) -> str:
    """Extract plain text from children list."""
    parts = []
    for c in children:
        parts.append(_extract_run_text(c))
    return "".join(parts)


# =============================================================================
# Block renderers
# =============================================================================


def _render_paragraph(doc: Document, block: dict[str, Any], styles: StyleConfig, md_path: Path = Path(".")) -> None:
    """Render a markdown paragraph with inline formatting."""
    para = doc.add_paragraph()
    apply_format_to_para(para, styles.body)

    children = block.get("children", [])
    for run_data in children:
        _render_inline_run(para, run_data, styles, md_path)


def _render_heading(doc: Document, block: dict[str, Any]) -> None:
    """Render a heading block."""
    level = block.get("level", 1)
    text = block.get("text", "")
    style_name = f"Heading {max(1, min(6, level))}"
    doc.add_paragraph(text, style=style_name)


def _render_table(doc: Document, block: dict[str, Any], table_style: TableStyle, md_path: Path = Path(".")) -> None:
    """Render a markdown table as an academic three-line table.

    Each cell may be a plain string (simple text) or a list of inline-run
    dicts (supports inline math, bold, italic, etc.).
    """
    header = block.get("header", [])
    rows = block.get("rows", [])

    if not header and not rows:
        return

    # Determine column count
    col_count = max(len(header), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return

    # Normalize header / rows to column count
    _norm: list = [""] * (col_count - len(header))
    header = list(header) + _norm
    norm_rows = [list(r) + [""] * (col_count - len(r)) for r in rows]

    table = doc.add_table(rows=1 + len(norm_rows), cols=col_count)
    table.autofit = True

    # Build a StyleConfig so _render_inline_run uses table fonts for math fallback
    _tbl_body = TextStyle(
        font_name=table_style.font_name,
        font_name_ascii=getattr(table_style, "font_name_ascii", table_style.font_name),
        font_size=table_style.font_size,
        line_spacing=table_style.line_spacing,
    )
    _tbl_styles = StyleConfig(body=_tbl_body)

    # Header row
    _fill_table_row(table.rows[0], header, table_style, _tbl_styles, md_path, is_header=True)

    # Data rows
    for i, row_data in enumerate(norm_rows):
        _fill_table_row(table.rows[i + 1], row_data, table_style, _tbl_styles, md_path, is_header=False)

    # Apply three-line formatting
    make_three_line_table(table)
    format_table_content(table, table_style)


def _fill_table_row(row, cell_data_list: list, table_style: TableStyle,
                    tbl_styles: StyleConfig, md_path: Path, *, is_header: bool) -> None:
    """Fill a table row, rendering inline runs (text, math, bold, etc.)."""
    for i, cell_data in enumerate(cell_data_list):
        cell = row.cells[i]
        # Clear default empty paragraph
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = table_style.line_spacing

        if isinstance(cell_data, str):
            # Plain text (backward-compatible fast path)
            run = para.add_run(cell_data.strip())
            apply_font_to_run(run, table_style)
            if is_header and table_style.header_bold:
                run.bold = True
        elif isinstance(cell_data, list):
            # Inline runs — supports inline_math, bold, italic, etc.
            for run_data in cell_data:
                _render_inline_run(para, run_data, tbl_styles, md_path)
            # Apply table font to non-math runs (math runs must keep Cambria Math)
            for run in para.runs:
                if _is_math_run(run):
                    continue
                apply_font_to_run(run, table_style)
                if is_header and table_style.header_bold:
                    run.bold = True
        else:
            # Fallback
            run = para.add_run(str(cell_data).strip())
            apply_font_to_run(run, table_style)
            if is_header and table_style.header_bold:
                run.bold = True


def _prepare_image(img_path: Path, max_w: float, max_h: float) -> tuple | None:
    """Load, re-encode, and size an image for docx insertion.

    Returns ``(stream, width, height)`` on success, ``None`` on failure.
    The image is re-encoded through Pillow to strip problematic EXIF/metadata
    that python-docx's internal JPEG parser cannot handle.

    SVG images are converted to PNG via svglib + reportlab before embedding
    since python-docx cannot render SVG natively.
    """
    from docx.shared import Inches

    try:
        from PIL import Image
        import io as _io
    except ImportError:
        return None

    is_svg = img_path.suffix.lower() == ".svg"

    # ---- Open image (SVG → PNG via svglib + reportlab) ----
    try:
        if is_svg:
            from svglib.svglib import svg2rlg
            from reportlab.graphics import renderPM

            drawing = svg2rlg(str(img_path))
            png_bytes = renderPM.drawToString(drawing, fmt="PNG")
            img = Image.open(_io.BytesIO(png_bytes))
        else:
            img = Image.open(img_path)
    except Exception:
        return None

    # ---- Size & re-encode ----
    try:
        max_width = Inches(max_w)
        max_height = Inches(max_h)
        aspect = img.width / img.height
        width = max_width
        height = Inches(width.inches / aspect)
        if height > max_height:
            height = max_height
            width = Inches(height.inches * aspect)

        image_stream = _io.BytesIO()
        save_format = "PNG" if is_svg else (img.format or "JPEG")
        if save_format.upper() in ("JPEG", "JPG") and img.mode == "RGBA":
            img = img.convert("RGB")
        img.save(image_stream, format=save_format)
    except (OSError, IOError, ValueError):
        return None

    image_stream.seek(0)
    return image_stream, width, height


def _render_image(doc: Document, block: dict[str, Any], md_path: Path) -> None:
    """Render an image block (full-width, centred)."""
    src = block.get("src", "")

    img_path = md_path.parent / src if not Path(src).is_absolute() else Path(src)
    if not img_path.exists():
        para = doc.add_paragraph()
        run = para.add_run(f"[图片缺失: {src}]")
        run.italic = True
        return

    result = _prepare_image(img_path, 5.5, 7.0)
    if result is None:
        para = doc.add_paragraph()
        run = para.add_run(f"[无法加载图片: {src}]")
        run.italic = True
        return

    image_stream, width, height = result
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(image_stream, width=width, height=height)


def _render_inline_image(para, src: str, alt: str, md_path: Path) -> None:
    """Render an image inline within a paragraph (smaller max size)."""
    img_path = md_path.parent / src if not Path(src).is_absolute() else Path(src)
    if not img_path.exists():
        run = para.add_run(f"[图片缺失: {src}]")
        run.font.italic = True
        return

    result = _prepare_image(img_path, 2.0, 2.0)
    if result is None:
        run = para.add_run(f"[无法加载图片: {src}]")
        run.font.italic = True
        return

    image_stream, width, height = result
    run = para.add_run()
    run.add_picture(image_stream, width=width, height=height)


def _render_list(doc: Document, block: dict[str, Any], styles: StyleConfig, md_path: Path = Path("."), level: int = 0) -> None:
    """Render an ordered or unordered list.

    List content uses the body text style for font/size/color so that list
    formatting follows the body style.  List-specific paragraph formatting
    (line_spacing, indent) is still applied from the list config.
    """
    list_style = styles.list
    ordered = block.get("ordered", False)
    items = block.get("items", [])

    for idx, item in enumerate(items):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = list_style.line_spacing
        para.paragraph_format.left_indent = Cm(list_style.indent * list_style.font_size * 0.0353 * (level + 1))

        # First-line indent for list items (matches body text indentation)
        first_indent_chars = getattr(list_style, "first_line_indent", 0)
        if first_indent_chars > 0:
            para.paragraph_format.first_line_indent = Pt(first_indent_chars * list_style.font_size)

        # Bullet or number prefix — same font as body text
        if ordered:
            prefix = f"{idx + 1}. "
        else:
            prefix = "• "

        prefix_run = para.add_run(prefix)
        apply_font_to_run(prefix_run, styles.body)

        # Render inline runs using body text style (preserves math, bold, italic, etc.)
        inline_runs = item.get("inline_runs", [])
        if inline_runs:
            for run_data in inline_runs:
                _render_inline_run(para, run_data, styles, md_path)
        else:
            # Fallback for plain text (backward compat)
            text = item.get("text", "")
            if text:
                run = para.add_run(text)
                apply_font_to_run(run, styles.body)

        # Handle nested lists
        for child in item.get("children", []):
            if isinstance(child, dict) and child.get("type") == "list":
                _render_list(doc, child, styles, md_path, level + 1)


def _render_code_block(doc: Document, block: dict[str, Any], code_style: CodeStyle) -> None:
    """Render a fenced code block as a bordered table box.

    Each line of code becomes a separate paragraph inside a single table cell.
    The table provides the visual "code frame" with borders on all four sides.
    """
    text = block.get("text", "")

    # Guard: skip empty code blocks
    if not text.strip():
        return

    lines = text.split("\n")
    if not lines:
        return

    # --- Create single-cell table ---
    table = doc.add_table(rows=1, cols=1)

    # Remove default table style (e.g. "Table Grid") to avoid border interference
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    for existing_style in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(existing_style)

    # --- Apply table-level borders ---
    if code_style.show_border:
        border_sz = int(code_style.border_width * 8)  # pt → eighths
        for position in ("top", "bottom", "left", "right"):
            _set_table_border(table, position, border_sz, code_style.border_color)
    else:
        for position in ("top", "bottom", "left", "right"):
            _set_table_border(table, position, 0, "auto")

    # No interior borders (single cell)
    _set_table_border(table, "insideH", 0, "auto")
    _set_table_border(table, "insideV", 0, "auto")

    # --- Fill the cell ---
    cell = table.cell(0, 0)

    # Remove the default empty paragraph — we add our own per-line paragraphs
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # --- Render each line as a separate paragraph ---
    for i, line in enumerate(lines):
        para = cell.add_paragraph()

        para.paragraph_format.line_spacing = code_style.line_spacing
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(0)

        if line:
            run = para.add_run(line)
            apply_font_to_run(run, code_style)
        else:
            # Preserve empty lines with a zero-width space to prevent collapsing
            run = para.add_run("​")
            apply_font_to_run(run, code_style)

    # --- Spacer after code block ---
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(3)
    spacer.paragraph_format.space_after = Pt(3)


def _render_blockquote(doc: Document, block: dict[str, Any], styles: StyleConfig, md_path: Path = Path(".")) -> None:
    """Render a blockquote."""
    children = block.get("children", [])
    for child in children:
        ctype = child.get("type", "")
        if ctype == "paragraph":
            para = doc.add_paragraph()
            set_para_left_border(para, styles.blockquote.border_left_color, 12)
            para.paragraph_format.left_indent = Cm(styles.blockquote.left_indent * 2.54)
            para.paragraph_format.line_spacing = styles.blockquote.line_spacing
            for run_data in child.get("children", []):
                _render_inline_run(para, run_data, styles, md_path)
            # Style runs
            for run in para.runs:
                apply_font_to_run(run, styles.blockquote)
        elif ctype == "heading":
            level = child.get("level", 1)
            hs = styles.headings.get(f"h{level}", styles.headings["h1"])
            para = doc.add_paragraph(child.get("text", ""))
            set_para_left_border(para, styles.blockquote.border_left_color, 12)
            para.paragraph_format.left_indent = Cm(styles.blockquote.left_indent * 2.54)
            for run in para.runs:
                apply_font_to_run(run, hs)
        else:
            _render_block(doc, child, styles, md_path)


# =============================================================================
# Page setup
# =============================================================================


def _setup_page(doc: Document, config: AppConfig) -> None:
    """Configure page size and margins."""
    for section in doc.sections:
        # Page size
        if config.page.size == "Letter":
            section.page_width = Cm(21.59)
            section.page_height = Cm(27.94)
        else:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # Margins
        section.top_margin = Cm(config.page.top_margin)
        section.bottom_margin = Cm(config.page.bottom_margin)
        section.left_margin = Cm(config.page.left_margin)
        section.right_margin = Cm(config.page.right_margin)


def _apply_page_number_font(run) -> None:
    """Apply 五号 font (10.5 pt) to a page-number field run."""
    set_run_font(
        run,
        cn_name="宋体",
        en_name="Times New Roman",
        size=SIZE_WUHAO,
        bold=False,
        italic=False,
    )


def _setup_page_number(doc: Document, page_number: PageNumberConfig) -> None:
    """Add page numbers to the footer of content sections.

    Page numbers are centred, use Arabic numerals (五号 Times New Roman), and
    appear only on sections that should be numbered.

    When *page_number.enabled* is False this is a no-op.

    Semantics of ``start_at``
    -------------------------
    ``start_at`` is the **physical page** of the document where Arabic
    numbering begins (with "1").  Pages before *start_at* carry no page
    number.  A section break inserted before the body content separates the
    front-matter section(s) from the numbered content section(s).

    * start_at = 1 → all pages numbered from 1 (single section)
    * start_at > 1 → first *(start_at - 1)* pages unnumbered;
      subsequent pages start at "1"
    """
    if not page_number.enabled:
        return

    sections = doc.sections
    if not sections:
        return

    # ---- Determine which sections get page numbers ----
    if page_number.start_at > 1 and len(sections) > 1:
        # Front-matter section(s) get no footer; numbering starts in
        # the first content section at 1.
        numbered_sections = sections[1:]  # skip cover/TOC section
        restart_at = 1
    else:
        # Single section or start_at == 1 — number everything.
        numbered_sections = list(sections)
        restart_at = page_number.start_at

    for i, section in enumerate(numbered_sections):
        footer = section.footer
        footer.is_linked_to_previous = False

        # --- Start with a clean paragraph (no Footer style) ---
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # --- Nuke the "Footer" paragraph style ---
        pPr = para._element.get_or_add_pPr()
        for stale in pPr.findall(qn("w:pStyle")):
            pPr.remove(stale)

        # --- Paragraph-level default run properties (五号) ---
        _para_rPr = OxmlElement("w:rPr")
        _rFonts = OxmlElement("w:rFonts")
        _rFonts.set(qn("w:ascii"), "Times New Roman")
        _rFonts.set(qn("w:hAnsi"), "Times New Roman")
        _rFonts.set(qn("w:eastAsia"), "宋体")
        _rFonts.set(qn("w:cs"), "Times New Roman")
        _para_rPr.append(_rFonts)
        _sz = OxmlElement("w:sz")
        _sz.set(qn("w:val"), "21")  # 10.5 pt = 五号
        _para_rPr.append(_sz)
        _szCs = OxmlElement("w:szCs")
        _szCs.set(qn("w:val"), "21")
        _para_rPr.append(_szCs)
        pPr.insert(0, _para_rPr)

        # --- Build PAGE field ---
        run_begin = para.add_run()
        _apply_page_number_font(run_begin)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin._element.append(fld_begin)

        run_instr = para.add_run()
        _apply_page_number_font(run_instr)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run_instr._element.append(instr)

        run_sep = para.add_run()
        _apply_page_number_font(run_sep)
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._element.append(fld_sep)

        run_text = para.add_run(str(restart_at))
        _apply_page_number_font(run_text)

        run_end = para.add_run()
        _apply_page_number_font(run_end)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._element.append(fld_end)

        # --- Restart page numbering for the first numbered section ---
        if i == 0:
            sect_pr = section._sectPr
            if sect_pr is not None:
                pg_num_type = OxmlElement("w:pgNumType")
                pg_num_type.set(qn("w:start"), str(restart_at))
                # Remove any existing pgNumType first
                for existing in sect_pr.findall(qn("w:pgNumType")):
                    sect_pr.remove(existing)
                sect_pr.append(pg_num_type)


# =============================================================================
# Cover page
# =============================================================================


def _build_cover(doc: Document, config: AppConfig) -> None:
    cover = config.cover

    # Report Title (with generous top spacing)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(72)
    run = title_p.add_run(cover.title)
    set_run_font(run, cn_name='黑体', en_name='Times New Roman',
                 size=Pt(28), bold=True)

    # Student info (first line spaced below title)
    lines = [
        ('姓    名：', cover.author),
        ('班    级：', cover.class_info),
        ('学    号：', cover.student_id),
        ('任课教师：', cover.teacher),
        ('实验日期：', cover.date),
    ]
    for i, (label, value) in enumerate(lines):
        p = add_cover_line(doc, label, value)
        if i == 0:
            p.paragraph_format.space_before = Pt(54)

    doc.add_page_break()

# =============================================================================
# TOC
# =============================================================================


def _build_toc(doc: Document, config: AppConfig) -> None:
    """Insert a Word Table of Contents field."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("目  录")
    set_run_font(run, cn_name="黑体", size=SIZE_SANHAO, bold=True)

    # TOC field
    insert_toc_field(doc, config.toc.level)

    # Page break after TOC
    doc.add_page_break()


# =============================================================================
# Block dispatcher
# =============================================================================


def _render_block(doc: Document, block: dict[str, Any], styles: StyleConfig, md_path: Path) -> None:
    """Dispatch a single block to the appropriate renderer."""
    btype = block.get("type", "")

    if btype == "heading":
        _render_heading(doc, block)

    elif btype == "paragraph":
        _render_paragraph(doc, block, styles, md_path)

    elif btype == "table":
        _render_table(doc, block, styles.table, md_path)

    elif btype == "image":
        _render_image(doc, block, md_path)

    elif btype == "list":
        _render_list(doc, block, styles, md_path)

    elif btype == "code_block":
        _render_code_block(doc, block, styles.code)

    elif btype == "blockquote":
        _render_blockquote(doc, block, styles, md_path)

    elif btype == "thematic_break":
        if styles.render_thematic_break:
            doc.add_paragraph("—" * 30)

    elif btype == "blank_line":
        pass  # skip — paragraph spacing handles separation

    elif btype == "block_math":
        latex = block.get("text", "")
        add_math_paragraph(doc, latex)


# =============================================================================
# Main entry point
# =============================================================================


def build_docx(
    ast: list[dict[str, Any]],
    config: AppConfig,
    md_path: str | Path,
) -> Document:
    """Build a python-docx Document from a parsed AST and configuration.

    Parameters
    ----------
    ast : list[dict]
        Parsed markdown block list from ``parse_markdown()``.
    config : AppConfig
        Full application configuration.
    md_path : str | Path
        Path to the source markdown file (for resolving image paths).

    Returns
    -------
    Document
        The built python-docx Document object (before saving).
    """
    md_path = Path(md_path)

    doc = Document()

    # --- 清除 python-docx 默认写入的文档属性 ---
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""

    # --- Page setup ---
    _setup_page(doc, config)

    # --- Normal style ---
    setup_normal_style(doc, config.styles.body)

    # --- Heading styles ---
    setup_heading_styles(doc, config.styles.headings)

    # --- TOC styles ---
    setup_toc_style(doc, config.toc)

    # --- Cover page ---
    if config.cover.enabled:
        _build_cover(doc, config)

    # --- Table of Contents ---
    if config.toc.enabled:
        _build_toc(doc, config)

    # --- Section break before body (when page numbering starts later) ---
    pn = config.page.page_number
    if pn.enabled and pn.start_at > 1:
        doc.add_section()
        # Re-apply page geometry to the new section (add_section copies from
        # the previous section, but _setup_page ensures consistency).
        _setup_page(doc, config)

    # --- Content ---
    for block in ast:
        _render_block(doc, block, config.styles, md_path)

    # --- Page number in footer (after all sections exist) ---
    _setup_page_number(doc, pn)

    return doc
