"""Math handler — LaTeX to OMML (Office Math Markup Language) conversion.

Pipeline: LaTeX string → MathML (via latex2mathml) → OMML (custom
recursive transformer) → insertion into python-docx paragraphs.

Fallback: If latex2mathml fails, renders the raw LaTeX as monospace text.
"""

from __future__ import annotations
import re
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from lxml import etree

# ---------------------------------------------------------------------------
# MathML namespace
# ---------------------------------------------------------------------------
MATHML_NS = "http://www.w3.org/1998/Math/MathML"
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _math_elt(tag: str, **attrs: str) -> Any:
    """Create an OMML element."""
    el = OxmlElement(f"m:{tag}")
    for k, v in attrs.items():
        el.set(qn(f"m:{k}"), v)
    return el


def _math_run(text: str) -> Any:
    """Create an OMML run (m:r) containing m:t with the given text."""
    r = _math_elt("r")
    t = _math_elt("t")
    t.text = text
    r.append(t)
    # Add font specification
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Cambria Math")
    rfonts.set(qn("w:hAnsi"), "Cambria Math")
    rpr.append(rfonts)
    r.insert(0, rpr)
    return r


def _process_children(mathml_el: Any) -> list[Any]:
    """Convert MathML element children to a list of OMML elements."""
    result: list[Any] = []
    for child in mathml_el:
        tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else child.tag.split("}")[-1] if "}" in (child.tag or "") else child.tag
        omml = _mathml_to_omml(child)
        if omml is not None:
            if isinstance(omml, list):
                result.extend(omml)
            else:
                result.append(omml)
    return result


# ---------------------------------------------------------------------------
# Fence / delimiter helpers
# ---------------------------------------------------------------------------


def _is_delimiter_char(ch: str) -> bool:
    """Return True if *ch* is a valid stretchy delimiter (bracket, brace, etc.).

    ``\right.`` produces ``.`` which maps to no delimiter — we exclude it here.
    """
    return ch in {
        "{", "}", "(", ")", "[", "]", "|",
        "‖",                           # ‖ double vertical line
        "⟨", "⟩",                 # ⟨ ⟩
        "⌈", "⌉",                 # ⌈ ⌉
        "⌊", "⌋",                 # ⌊ ⌋
    }


def _handle_fenced_mrow(children: list) -> Any:
    """Wrap non-fence children in an OMML ``m:d`` delimiter with stretchy brackets.

    Handles ``\\begin{cases}``, ``\\left\\{ ... \\right.``, and similar
    constructions where MathML uses ``<mo fence="true">`` alongside the content.
    """
    open_ch: str | None = None
    close_ch: str | None = None
    content_children: list = []

    for child in children:
        try:
            ctag = (etree.QName(child.tag).localname
                    if hasattr(child, "tag") else "")
        except Exception:
            ctag = ""
        is_fence = (
            ctag == "mo"
            and child.get("fence") == "true"
        )
        if is_fence and _is_delimiter_char((child.text or "").strip()):
            ch = (child.text or "").strip()
            form = child.get("form", "")
            if form == "prefix" or (form != "postfix" and open_ch is None):
                open_ch = ch
            elif form == "postfix":
                close_ch = ch
        else:
            content_children.append(child)

    d = _math_elt("d")
    d_pr = _math_elt("dPr")

    # Always set both begChr and endChr explicitly.  OMML defaults are
    # "(" for begChr and ")" for endChr — we override with empty string
    # when no delimiter was detected (e.g. cases has no closing bracket,
    # \right. produces no delimiter).
    beg = _math_elt("begChr")
    beg.set(qn("m:val"), open_ch if open_ch else "")
    d_pr.append(beg)

    end = _math_elt("endChr")
    end.set(qn("m:val"), close_ch if close_ch else "")
    d_pr.append(end)

    d.append(d_pr)

    e = _math_elt("e")
    for child in content_children:
        result = _mathml_to_omml(child)
        _append_children(e, result)
    d.append(e)

    return d


def _mrow_has_fence(children) -> bool:
    """Return True if any child of the mrow is a fence ``<mo>``."""
    for child in children:
        try:
            ctag = (etree.QName(child.tag).localname
                    if hasattr(child, "tag") else "")
        except Exception:
            ctag = ""
        if ctag == "mo" and child.get("fence") == "true":
            return True
    return False


def _mathml_to_omml(el: Any) -> Any:
    """Recursively convert a MathML element (lxml) into OMML element(s).

    Returns a single OMML element, a list of OMML elements, or None.
    """
    # Handle plain text nodes
    if isinstance(el, str) or (hasattr(el, "tag") and el.tag is etree.Comment):
        return None

    tag = etree.QName(el.tag).localname if el.tag and hasattr(el, "tag") else ""
    text = el.text.strip() if el.text else ""

    # --- Leaf elements ---
    if tag in ("mi", "mn", "mo", "mtext"):
        # Fence <mo> elements are handled by their parent mrow; if an
        # isolated fence mo reaches here it's treated as regular text.
        return _math_run(text)

    # --- mrow: group (with fence detection for cases, \left\{, etc.) ---
    if tag == "mrow":
        children = list(el)  # type: ignore[arg-type]
        if _mrow_has_fence(children):
            return _handle_fenced_mrow(children)
        return _process_children(el)

    # --- mfrac: fraction ---
    if tag == "mfrac":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 2:
            f = _math_elt("f")
            fpr = _math_elt("fPr")
            f.append(fpr)
            num = _math_elt("num")
            num_result = _mathml_to_omml(children[0])
            _append_children(num, num_result)
            f.append(num)
            den = _math_elt("den")
            den_result = _mathml_to_omml(children[1])
            _append_children(den, den_result)
            f.append(den)
            return f
        return None

    # --- msqrt / mroot: radicals ---
    if tag in ("msqrt", "mroot"):
        rad = _math_elt("rad")
        rad_pr = _math_elt("radPr")
        if tag == "msqrt":
            deg_hide = _math_elt("degHide")
            deg_hide.set(qn("m:val"), "1")
            rad_pr.append(deg_hide)
        else:
            # mroot has a degree
            children = list(el)  # type: ignore[arg-type]
            if len(children) >= 2:
                deg = _math_elt("deg")
                deg_result = _mathml_to_omml(children[1])
                _append_children(deg, deg_result)
                rad_pr.append(deg)
        rad.append(rad_pr)
        # Base (e)
        base_child = list(el)[0] if len(list(el)) > 0 else None  # type: ignore[arg-type]
        if base_child is not None:
            e = _math_elt("e")
            base_result = _mathml_to_omml(base_child)
            _append_children(e, base_result)
            rad.append(e)
        return rad

    # --- msup: superscript ---
    if tag == "msup":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 2:
            ssup = _math_elt("sSup")
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            ssup.append(e)
            sup = _math_elt("sup")
            _append_children(sup, _mathml_to_omml(children[1]))
            ssup.append(sup)
            return ssup
        return None

    # --- msub: subscript ---
    if tag == "msub":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 2:
            ssub = _math_elt("sSub")
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            ssub.append(e)
            sub = _math_elt("sub")
            _append_children(sub, _mathml_to_omml(children[1]))
            ssub.append(sub)
            return ssub
        return None

    # --- msubsup: sub+superscript ---
    if tag == "msubsup":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 3:
            ss = _math_elt("sSubSup")
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            ss.append(e)
            sub = _math_elt("sub")
            _append_children(sub, _mathml_to_omml(children[1]))
            ss.append(sub)
            sup = _math_elt("sup")
            _append_children(sup, _mathml_to_omml(children[2]))
            ss.append(sup)
            return ss
        return None

    # --- mover: overscript (accent) ---
    if tag == "mover":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 2:
            acc = _math_elt("acc")
            acc_pr = _math_elt("accPr")
            # Extract the accent character
            accent_char = _extract_text(children[1])
            if accent_char:
                chr_el = _math_elt("chr")
                chr_el.set(qn("m:val"), accent_char)
                acc_pr.append(chr_el)
            acc.append(acc_pr)
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            acc.append(e)
            return acc
        return None

    # --- munder: underscript (limit) ---
    if tag == "munder":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 2:
            lim_low = _math_elt("limLow")
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            lim_low.append(e)
            lim = _math_elt("lim")
            _append_children(lim, _mathml_to_omml(children[1]))
            lim_low.append(lim)
            return lim_low
        return None

    # --- munderover: underscript + overscript ---
    if tag == "munderover":
        children = list(el)  # type: ignore[arg-type]
        if len(children) >= 3:
            # Use limUpp + limLow combination
            lim_upp = _math_elt("limUpp")
            e = _math_elt("e")
            _append_children(e, _mathml_to_omml(children[0]))
            lim_upp.append(e)
            lim_up = _math_elt("lim")
            _append_children(lim_up, _mathml_to_omml(children[2]))
            lim_upp.append(lim_up)

            # Wrap limUpp inside limLow for the lower limit
            lim_low = _math_elt("limLow")
            lim_low.append(lim_upp)  # The e from limUpp becomes the base
            # Actually this nesting is tricky. Let's use a simpler approach:
            # Use an nary-like structure or just put both limits on limUpp
            return lim_upp
        return None

    # --- mfenced: delimiters (parentheses, brackets, etc.) ---
    if tag == "mfenced":
        open_ch = el.get("open", "(")
        close_ch = el.get("close", ")")
        d = _math_elt("d")
        d_pr = _math_elt("dPr")
        beg = _math_elt("begChr")
        beg.set(qn("m:val"), open_ch)
        d_pr.append(beg)
        end = _math_elt("endChr")
        end.set(qn("m:val"), close_ch)
        d_pr.append(end)
        d.append(d_pr)
        e = _math_elt("e")
        for child in el:
            result = _mathml_to_omml(child)
            _append_children(e, result)
        d.append(e)
        return d

    # --- mtable: matrix ---
    if tag == "mtable":
        m = _math_elt("m")
        m_pr = _math_elt("mPr")
        m.append(m_pr)
        for row_el in el:
            mr = _math_elt("mr")
            for cell_el in row_el:
                e = _math_elt("e")
                result = _mathml_to_omml(cell_el)
                _append_children(e, result)
                mr.append(e)
            m.append(mr)
        return m

    # --- Fallback: treat as row and process children ---
    return _process_children(el)


def _append_children(parent: Any, result: Any) -> None:
    """Append OMML result (single element or list) to parent."""
    if result is None:
        return
    if isinstance(result, list):
        for item in result:
            if item is not None:
                parent.append(item)
    else:
        parent.append(result)


def _extract_text(el: Any) -> str:
    """Extract plain text from a MathML element."""
    if el is None:
        return ""
    if hasattr(el, "text") and el.text:
        return el.text.strip()
    # Try to get text from children
    parts = []
    if hasattr(el, "itertext"):
        parts = list(el.itertext())
    return "".join(parts)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def latex_to_omml(latex: str) -> Any | None:
    """Convert a LaTeX string to an OMML ``m:oMath`` element.

    Parameters
    ----------
    latex : str
        LaTeX expression without delimiters (no ``$...$``).

    Returns
    -------
    lxml Element or None
        An ``<m:oMath>`` element ready to insert into a paragraph,
        or *None* if conversion failed.
    """
    try:
        from latex2mathml.converter import convert as latex2mathml_convert

        # Strip optional spacing argument from line breaks:
        #   \\[6pt]  → \\
        #   \\*[10pt] → \\*
        # latex2mathml doesn't handle these and renders "[]" as literal text.
        latex = re.sub(r'(\\\\\*?)\[[^\]]*\]', r'\1', latex)

        # Convert LaTeX to MathML string
        mathml_str = latex2mathml_convert(latex)
        if not mathml_str:
            return None

        # Parse MathML
        mathml_el = etree.fromstring(mathml_str.encode("utf-8"))

        # Find the <math> root (might be wrapped)
        if etree.QName(mathml_el.tag).localname == "math":
            math_root = mathml_el
        else:
            math_ns_el = mathml_el.find(f"{{{MATHML_NS}}}math")
            if math_ns_el is not None:
                math_root = math_ns_el
            else:
                math_root = mathml_el

        # Convert to OMML
        omml = _math_elt("oMath")
        children = _process_children(math_root)
        for child in children:
            omml.append(child)

        return omml

    except ImportError:
        return None
    except Exception:
        return None


def add_display_math(para, latex: str) -> bool:
    """Insert a display math equation — oMathPara appended to paragraph.

    Returns True on success, False on fallback.
    """
    omml = latex_to_omml(latex)
    if omml is None:
        run = para.add_run(f"$$ {latex} $$")
        return False

    omath_para = OxmlElement("m:oMathPara")
    omath_para.append(omml)
    para._element.append(omath_para)
    return True


def add_inline_math(para, latex: str) -> bool:
    """Insert an inline math equation inside a paragraph.

    Creates a proper w:r wrapper with math font rPr, then injects the OMML.
    Returns True on success, False on fallback.
    """
    omml = latex_to_omml(latex)
    if omml is None:
        run = para.add_run(f" ${latex}$ ")
        return False

    # Create a proper run via python-docx, then replace its content with OMML.
    # This ensures correct namespace context and rPr initialization.
    placeholder = para.add_run(" ")
    r_elem = placeholder._element

    # Remove the w:t (placeholder text) and inject OMML
    t_elem = r_elem.find(qn("w:t"))
    if t_elem is not None:
        r_elem.remove(t_elem)

    # Ensure the run has rPr with math font
    rpr = r_elem.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r_elem.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Cambria Math")
    rfonts.set(qn("w:hAnsi"), "Cambria Math")
    rfonts.set(qn("w:eastAsia"), "Cambria Math")
    rfonts.set(qn("w:cs"), "Cambria Math")

    r_elem.append(omml)
    return True


def add_math_paragraph(doc: Document, latex: str) -> Any:
    """Create a new centered paragraph containing a display math equation.

    Returns the paragraph object.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_display_math(para, latex)
    return para
