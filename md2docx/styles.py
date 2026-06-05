"""Word style helpers — clean, readable API for docx formatting.

Inspired by the direct, self-documenting style of standalone report scripts.
All font sizes follow the Chinese academic convention (中文字号).
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# =============================================================================
# Chinese font size constants (中文字号 → pt)
# =============================================================================
SIZE_CHUHAO = Pt(42)       # 初号
SIZE_XIAOCHU = Pt(36)      # 小初
SIZE_YIHAO = Pt(26)        # 一号
SIZE_XIAOYI = Pt(24)       # 小一
SIZE_ERHAO = Pt(22)        # 二号
SIZE_XIAOER = Pt(18)       # 小二
SIZE_SANHAO = Pt(16)       # 三号
SIZE_XIAOSAN = Pt(15)      # 小三
SIZE_SIHAO = Pt(14)        # 四号
SIZE_XIAOSI = Pt(12)       # 小四（正文）
SIZE_WUHAO = Pt(10.5)      # 五号（表格/脚注）
SIZE_XIAOWU = Pt(9)        # 小五

# =============================================================================
# Alignment map
# =============================================================================
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# =============================================================================
# Run (character) formatting
# =============================================================================


def set_run_font(
    run,
    cn_name: str = "宋体",
    en_name: str = "Times New Roman",
    size: Pt = SIZE_XIAOSI,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """Apply font properties to a run — one call, self-documenting.

    Usage::

        set_run_font(run, cn_name="黑体", size=SIZE_SANHAO, bold=True)
    """
    # Set ASCII/HAnsi via python-docx API
    run.font.name = en_name
    run.font.size = size
    run.bold = bold
    run.italic = italic

    # set East-Asian font explicitly
    rpr = run._element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    rfonts.set(qn("w:eastAsia"), cn_name)

    # Color
    if color is not None:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Default to black if color is None


# =============================================================================
# Paragraph formatting
# =============================================================================


def set_para_format(
    para,
    line_spacing: float = 1.5,
    alignment: str = "left",
    first_line_indent_chars: int = 0,
    font_size_for_indent: Pt | None = None,
    space_before_pt: int = 0,
    space_after_pt: int = 0,
) -> None:
    """Apply paragraph-level formatting — one call, self-documenting.

    Usage::

        set_para_format(para, line_spacing=1.5, alignment="justify",
                        first_line_indent_chars=2, font_size_for_indent=SIZE_XIAOSI)
    """
    para.paragraph_format.line_spacing = line_spacing

    if alignment in _ALIGN:
        para.alignment = _ALIGN[alignment]

    if first_line_indent_chars > 0 and font_size_for_indent is not None:
        # 1 Chinese char ≈ font_size in pt → cm
        indent_cm = first_line_indent_chars * font_size_for_indent.pt * 0.0353
        para.paragraph_format.first_line_indent = Cm(indent_cm)

    if space_before_pt is not None:
        para.paragraph_format.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        para.paragraph_format.space_after = Pt(space_after_pt)


# =============================================================================
# Paragraph shading / borders (code blocks, blockquotes)
# =============================================================================


def set_para_shading(para, hex_color: str) -> None:
    """Set paragraph background color (e.g. for code blocks)."""
    ppr = para._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    ppr.append(shd)


def set_para_left_border(para, hex_color: str = "999999", width_eighths_pt: int = 12) -> None:
    """Add a left border line to a paragraph (e.g. for blockquotes)."""
    ppr = para._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_eighths_pt))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    p_bdr.append(left)
    ppr.append(p_bdr)


# =============================================================================
# Apply a style-config object to a run / paragraph
# =============================================================================


def apply_font_to_run(run, style) -> None:
    """Apply font properties from a TextStyle/HeadingStyle/TableStyle config to a run."""
    cn = style.font_name
    en = getattr(style, "font_name_ascii", cn)
    color = None
    if getattr(style, "color", None):
        try:
            color = RGBColor.from_string(style.color)
        except Exception:
            pass
    set_run_font(
        run,
        cn_name=cn,
        en_name=en,
        size=Pt(style.font_size),
        bold=getattr(style, "bold", False),
        italic=getattr(style, "italic", False),
        color=color,
    )


def apply_format_to_para(para, style) -> None:
    """Apply paragraph formatting from a TextStyle config (spacing, indent, alignment)."""
    set_para_format(
        para,
        line_spacing=getattr(style, "line_spacing", 1.5),
        alignment=getattr(style, "alignment", "left"),
        first_line_indent_chars=getattr(style, "first_line_indent", 0),
        font_size_for_indent=Pt(getattr(style, "font_size", 12)),
        space_before_pt=getattr(style, "space_before", 0),
        space_after_pt=getattr(style, "space_after", 0),
    )


# =============================================================================
# Cover page helpers
# =============================================================================


def add_cover_line(doc, text, value='__________', font_size=Pt(16), font_name="宋体"):
    """Add a line on the cover page."""
    p = doc.add_paragraph()
    p.alignment =  _ALIGN["center"]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(text)
    run_label.font.name = font_name
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run_label.font.size = font_size
    run_label.bold = True
    run_value = p.add_run('  ' + value)
    run_value.font.name = font_name
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run_value.font.size = font_size
    run_value.underline = True
    return p


# =============================================================================
# Table of Contents
# =============================================================================


def insert_toc_field(doc: Document, levels: int = 3) -> None:
    """Insert a Word TOC field (requires right-click → Update Field in Word)."""
    para = doc.add_paragraph()

    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    run2._element.append(instr)

    run3 = para.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_sep)

    hint = para.add_run("（请在 Word 中右键此处 → 更新域，以生成目录）")
    set_run_font(hint, cn_name="宋体", size=SIZE_XIAOSI, italic=True,
                 color=RGBColor(128, 128, 128))

    run4 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run4._element.append(fld_end)


def setup_toc_style(doc: Document, toc_style) -> None:
    """Create/override Word built-in TOC 1~9 styles with configured font, size, color.

    When the user updates the TOC field in Word, the generated entries inherit
    these style definitions.  A single config object is applied to all TOC levels
    by default (Chinese academic convention — all TOC lines same font/size).

    The ``w:customStyle`` attribute is explicitly removed after creation so that
    Word recognises these as the real built-in TOC styles rather than ignoring
    them during TOC field updates.
    """
    from docx.enum.style import WD_STYLE_TYPE

    cn = getattr(toc_style, "font_name", "宋体")
    en = getattr(toc_style, "font_name_ascii", "Times New Roman")

    for i in range(1, 10):  # TOC 1 … TOC 9
        style_name = f"TOC {i}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            # Style doesn't exist in the default template — create it
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Remove customStyle flag so Word treats this as a built-in TOC style
        style_elem = style.element
        style_elem.attrib.pop(qn("w:customStyle"), None)

        _set_style_fonts(style, cn, en)
        style.font.size = Pt(getattr(toc_style, "font_size", 12))
        style.font.bold = getattr(toc_style, "bold", True)
        style.font.italic = getattr(toc_style, "italic", False)

        color_value = getattr(toc_style, "color", None)
        if color_value:
            style.font.color.rgb = RGBColor.from_string(color_value)

        # Line spacing — compact single spacing
        style.paragraph_format.line_spacing = 1.0

        # Per-level left indent (matches standard Word TOC behaviour:
        # TOC 1 flush left, TOC 2 indented ~0.75 cm, TOC 3 ~1.5 cm, etc.)
        if i > 1:
            style.paragraph_format.left_indent = Cm((i - 1) * 0.75)


# =============================================================================
# Style setup (Normal + Heading 1-6)
# =============================================================================


def _set_style_fonts(style, cn_font: str, en_font: str) -> None:
    """Set both ASCII/HAnsi and East-Asian fonts for a style, clear theme references."""
    style.font.name = cn_font

    rpr = style.element.rPr
    rFonts = rpr.find(qn('w:rFonts'))

    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), en_font)

    # MUST clear theme font references
    for attr in (qn('w:asciiTheme'), qn('w:hAnsiTheme'),
                 qn('w:eastAsiaTheme'), qn('w:cstheme')):
        try:
            del rFonts.attrib[attr]
        except KeyError:
            pass


def setup_normal_style(doc: Document, body_style) -> None:
    """Override the Normal style with configured body font."""
    style = doc.styles["Normal"]
    
    cn_font = getattr(body_style, "font_name", "宋体")
    en_font = getattr(body_style, "font_name_ascii", "Times New Roman")
    
    _set_style_fonts(style, cn_font, en_font)
    style.font.size = Pt(getattr(body_style, "font_size", 12))
    style.font.bold = getattr(body_style, "bold", False)
    style.font.italic = getattr(body_style, "italic", False)
    color_value = getattr(body_style, "color", None)
    if color_value:
        style.font.color.rgb = RGBColor.from_string(color_value)
    else:
        style.font.color.rgb = RGBColor(0, 0, 0)  # Default to black if color is None
    
    pf = style.paragraph_format
    pf.line_spacing = getattr(body_style, "line_spacing", 1.5)
    pf.space_before = Pt(getattr(body_style, "space_before", 0))
    pf.space_after = Pt(getattr(body_style, "space_after", 0))
    align = getattr(body_style, "alignment", "justify")
    if align in _ALIGN:
        pf.alignment = _ALIGN[align]


def setup_heading_styles(doc: Document, heading_config: dict) -> None:
    """Override Word heading styles 1–6 — clear themes, set 黑体 + explicit black.

    heading_config maps 'h1'..'h6' to HeadingStyle objects.
    """
    for i, key in enumerate(("h1", "h2", "h3", "h4", "h5", "h6"), 1):
        hs = heading_config.get(key)
        if hs is None:
            continue

        style_name = f"Heading {i}"
        word_style = doc.styles[style_name]

        # Font
        cn = getattr(hs, "font_name", "黑体")
        en = getattr(hs, "font_name_ascii", "Times New Roman")
        _set_style_fonts(word_style, cn, en)

        # Size / bold / italic
        word_style.font.size = Pt(getattr(hs, "font_size", 14))
        word_style.font.bold = getattr(hs, "bold", True)
        word_style.font.italic = getattr(hs, "italic", False)

        # Color
        color_value = getattr(hs, "color", None)
        if color_value:
            word_style.font.color.rgb = RGBColor.from_string(color_value)
        else:
            word_style.font.color.rgb = RGBColor(0, 0, 0)  # Default to black if color is None


        # Paragraph format
        pf = word_style.paragraph_format
        align = getattr(hs, "alignment", "left")
        if align in _ALIGN:
            pf.alignment = _ALIGN[align]
        pf.space_before = Pt(getattr(hs, "space_before", 6))
        pf.space_after = Pt(getattr(hs, "space_after", 6))


# =============================================================================
# Academic three-line table
# =============================================================================


def make_three_line_table(table, header_rows: int = 1) -> None:
    """Format a python-docx Table as an academic three-line table.

    Border precedence in OOXML (highest to lowest):
      1. Cell-level exceptions (w:tcBorders) — overrides everything
      2. Table-level exceptions (w:tblBorders) — overrides style
      3. Table style conditional formatting
      4. Table style
      5. Document defaults

    Setting cell-level borders to 0 pt white would BLOCK table-level borders,
    so we work at the table level only (except for the header-underline, which
    is genuinely a cell-level detail).

    - Top border: 1.5 pt (12 eighths of a point)
    - Below header: 0.75 pt (6 eighths)
    - Bottom border: 1.5 pt (12 eighths)
    - No vertical or interior horizontal borders
    """

    # ---- 1. Remove default table style so it doesn't interfere ----
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    # Clear any inherited style (e.g. "Table Grid") that may inject borders
    for existing_style in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(existing_style)

    # ---- 2. Set table-level borders ----
    _set_table_border(table, "top", 12)          # 1.5 pt
    _set_table_border(table, "bottom", 12)       # 1.5 pt
    _set_table_border(table, "left", 0, "auto")         # no left
    _set_table_border(table, "right", 0, "auto")        # no right
    _set_table_border(table, "insideH", 0, "auto")      # no interior horizontal
    _set_table_border(table, "insideV", 0, "auto")      # no interior vertical

    # ---- 3. Cell-level: underline below header row(s) ----
    for i in range(header_rows):
        for cell in table.rows[i].cells:
            _set_cell_border(cell, "bottom", 6)  # 0.75 pt under header


def _is_math_run(run) -> bool:
    """Return True if *run* is an OMML math run (contains m:oMath child)."""
    for child in run._element:
        if child.tag.endswith("}oMath") or child.tag.endswith("}oMathPara"):
            return True
    return False


def format_table_content(table, table_style, header_rows: int = 1) -> None:
    """Apply font, size, alignment to all cells in a table.

    Math runs (OMML) are skipped — they must keep their Cambria Math font.
    """
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = _ALIGN.get(
                    getattr(table_style, "alignment", "center"),
                    WD_ALIGN_PARAGRAPH.CENTER,
                )
                para.paragraph_format.line_spacing = getattr(table_style, "line_spacing", 1.0)
                for run in para.runs:
                    if _is_math_run(run):
                        continue
                    apply_font_to_run(run, table_style)
                    if row_idx < header_rows and getattr(table_style, "header_bold", True):
                        run.bold = True


# =============================================================================
# Internal: XML-level table border helpers
# =============================================================================


def _set_cell_border(cell, position: str, sz: int, color: str = "000000") -> None:
    """Set a cell-level border (highest OOXML precedence — use sparingly).

    Parameters
    ----------
    sz : int
        Border width in eighths of a point.  Use 0 for no border.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    val = "none" if sz == 0 else "single"
    elem = OxmlElement(f"w:{position}")
    elem.set(qn("w:val"), val)
    elem.set(qn("w:sz"), str(sz))
    elem.set(qn("w:space"), "0")
    elem.set(qn("w:color"), color)

    existing = borders.find(qn(f"w:{position}"))
    if existing is not None:
        borders.remove(existing)
    borders.append(elem)


def _set_table_border(table, position: str, sz: int, color: str = "000000") -> None:
    """Set a table-level border.

    Parameters
    ----------
    sz : int
        Border width in eighths of a point (1 pt = 8).  Use 0 for no border.
    color : str
        Hex colour.  Use ``"auto"`` to let Word decide (typically black for
        visible borders, invisible for none/nil borders).
    """
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        # tblBorders should come before tblLook in tblPr
        tbl_look = tbl_pr.find(qn("w:tblLook"))
        if tbl_look is not None:
            tbl_pr.insert(list(tbl_pr).index(tbl_look), borders)
        else:
            tbl_pr.append(borders)

    val = "none" if sz == 0 else "single"
    elem = OxmlElement(f"w:{position}")
    elem.set(qn("w:val"), val)
    elem.set(qn("w:sz"), str(sz))
    elem.set(qn("w:space"), "0")
    elem.set(qn("w:color"), color)

    existing = borders.find(qn(f"w:{position}"))
    if existing is not None:
        borders.remove(existing)
    borders.append(elem)


